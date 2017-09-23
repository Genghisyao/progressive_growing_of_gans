# -*- coding:utf-8 -*-
from db_init import *
from crawl import *
# from docx import Document
# from docx.shared import Inches


def heat_calculate(lis):
    """
    计算热度————连续时间段加权平均值
    :param lis: i[0]=总评论，i[1]=总阅读, i[2]=总转载
    :return: heat_total=热度
    """
    heat_total = 0.0
    for i in lis:
        reply_total = float(i[0]) * 0.45
        read_total = float(i[1]) * 0.236
        forward_total = float(i[2]) * 0.314
        heat_total += (reply_total + read_total + forward_total)
    return heat_total/len(lis)

class Processing():
    """
    分析处理模块
    """

    def __init__(self, event_id, add_datetime, start_datetime, keywords, non_keywords):
        self.event_id = event_id  #事件id
        self.add_datetime = str(add_datetime)
        self.start_datetime = str(start_datetime)
        self.keywords = keywords  #完整关键词
        self.non_keywords = non_keywords  #完整反关键词
        keys = keywords.split(',')
        k_list = []
        for k in keys:
            k_list.append('%' + k + '%')
        self.key_list = k_list

    def newsSelect(self):
        """
        筛选————新闻
        """
        key = self.key_list
        sql_t = []
        for k in key:
            sql_t.append(" Ttitle LIKE '%s' OR Tcontent LIKE '%s'" % (k, k))
        sql_query = """SELECT * FROM news_text WHERE (""" + 'OR'.join(sql_t) + ") AND Tpublish_datetime > '%s'" % self.start_datetime
        query = session.execute(sql_query)
        sql = []
        for n in query:
            Tid = n.id
            event_id = self.event_id
            title = n.Ttitle.replace(':','\:')
            add_datetime = n.add_datetime
            publish_datetime = n.Tpublish_datetime
            author_id = ''
            author_name = n.Tauthor_name
            content = n.Tcontent
            reply_count = n.Treply
            read_count = 0
            like_count = 0
            collect_count = 0
            forward_count = 0
            channeltype = 'news'
            channel = n.site
            correlation = ''
            heatwords = ''
            if reply_count < 0:
                reply_count = 0
            if read_count < 0:
                read_count = 0
            if like_count < 0:
                like_count = 0
            if collect_count < 0:
                collect_count = 0
            if forward_count < 0:
                forward_count = 0
            weight_reply = 0.45  #评论权重
            weight_read = 0.236  #阅读权重
            weight_like = 0.0  #点赞权重
            weight_collect = 0.0  #收藏权重
            weight_forward = 0.314  #转载权重
            heat = reply_count * weight_reply + read_count * weight_read + like_count * weight_like + collect_count * weight_collect + forward_count * weight_forward
            sql.append(" ('%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s') " %
                       (Tid,event_id,title,add_datetime,publish_datetime,author_id,author_name,content,
                        reply_count,read_count,like_count,collect_count,forward_count,channeltype,channel,
                        correlation,heatwords,heat))
        if sql:
            sql_insert = """INSERT INTO text_select VALUES""" + ','.join(sql) \
                         + """ON DUPLICATE KEY UPDATE reply_count = VALUES(reply_count), read_count = VALUES(read_count), like_count = VALUES(like_count),collect_count = VALUES(collect_count),forward_count = VALUES(forward_count)"""
            session.execute(sql_insert)
            session.commit()
            # sql_insert = """
            #     INSERT INTO text_select (Tid,event_id,title,add_datetime,publish_datetime,author_id,author_name,content,
            #     reply_count,read_count,like_count,collect_count,forward_count,channeltype,channel,correlation,heatwords,heat)
            #     VALUES ('%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s')
            #     ON DUPLICATE KEY UPDATE reply_count = VALUES(reply_count)
            # """ % (Tid,event_id,title,add_datetime,publish_datetime,author_id,author_name,content,reply_count,read_count,
            #        like_count,collect_count,forward_count,channeltype,channel,correlation,heatwords,heat)
            # session.execute(sql_insert)
            # session.commit()

    def luntanSelect(self):
        """
        筛选————论坛
        """
        key = self.key_list
        sql_t = []
        for k in key:
            sql_t.append(" Ttitle LIKE '%s' OR Tcontent LIKE '%s'" % (k, k))
        sql_query = """SELECT * FROM luntan_text WHERE (""" + 'OR'.join(sql_t) + ") AND Tpublish_datetime > '%s'" % self.start_datetime
        query = session.execute(sql_query)
        sql = []
        # q_news = session.query(NewsText).filter(or_(NewsText.Ttitle.like(key), NewsText.Tcontent.like(key))).all()
        for n in query:
            Tid = n.id
            event_id = self.event_id
            title = n.Ttitle
            add_datetime = n.add_datetime
            publish_datetime = n.Tpublish_datetime
            author_id = n.Tauthor_id
            author_name = n.Tauthor_name
            content = n.Tcontent
            reply_count = n.Treply
            read_count = n.Tread
            like_count = 0
            collect_count = 0
            forward_count = 0
            channeltype = 'luntan'
            channel = 'bbs.tianya.cn'
            correlation = ''
            heatwords = ''
            if reply_count < 0:
                reply_count = 0
            if read_count < 0:
                read_count = 0
            if like_count < 0:
                like_count = 0
            if collect_count < 0:
                collect_count = 0
            if forward_count < 0:
                forward_count = 0
            weight_reply = 0.45  # 评论权重
            weight_read = 0.236  # 阅读权重
            weight_like = 0.0  # 点赞权重
            weight_collect = 0.0  # 收藏权重
            weight_forward = 0.314  # 转载权重
            heat = reply_count * weight_reply + read_count * weight_read + like_count * weight_like + collect_count * weight_collect + forward_count * weight_forward
            sql.append(" ('%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s') " %
                       (Tid, event_id, title, add_datetime, publish_datetime, author_id, author_name, content,
                        reply_count, read_count, like_count, collect_count, forward_count, channeltype, channel,
                        correlation, heatwords, heat))
        if sql:
            sql_insert = """INSERT INTO text_select VALUES""" + ','.join(sql) \
                         + """ON DUPLICATE KEY UPDATE reply_count = VALUES(reply_count), read_count = VALUES(read_count), like_count = VALUES(like_count),collect_count = VALUES(collect_count),forward_count = VALUES(forward_count)"""
            session.execute(sql_insert)
            session.commit()
            # sql_insert = """
            #     INSERT INTO text_select (Tid,event_id,title,add_datetime,publish_datetime,author_id,author_name,content,
            #     reply_count,read_count,like_count,collect_count,forward_count,channeltype,channel,correlation,heatwords,heat)
            #     VALUES ('%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s')
            #     ON DUPLICATE KEY UPDATE reply_count = VALUES(reply_count)
            # """ % (Tid, event_id, title, add_datetime, publish_datetime, author_id, author_name, content, reply_count,
            # read_count, like_count, collect_count, forward_count, channeltype, channel, correlation, heatwords, heat)
            # session.execute(sql_insert)
            # session.commit()

    def tiebaSelect(self):
        """
        筛选————贴吧
        """
        key = self.key_list
        sql_t = []
        for k in key:
            sql_t.append(" Ttitle LIKE '%s' OR Tcontent LIKE '%s'" % (k, k))
        sql_query = """SELECT * FROM tieba_text WHERE (""" + 'OR'.join(sql_t) + ") AND Tpublish_datetime > '%s'" % self.start_datetime
        query = session.execute(sql_query)
        sql = []
        # q_news = session.query(NewsText).filter(or_(NewsText.Ttitle.like(key), NewsText.Tcontent.like(key))).all()
        for n in query:
            Tid = n.id
            event_id = self.event_id
            title = n.Ttitle
            add_datetime = n.add_datetime
            publish_datetime = n.Tpublish_datetime
            author_id = n.Tauthor_id
            author_name = n.Tauthor_name
            content = n.Tcontent
            reply_count = n.Treply
            read_count = 0
            like_count = 0
            collect_count = 0
            forward_count = 0
            channeltype = 'tieba'
            channel = 'tieba.baidu.com'
            correlation = ''
            heatwords = ''
            if reply_count < 0:
                reply_count = 0
            if read_count < 0:
                read_count = 0
            if like_count < 0:
                like_count = 0
            if collect_count < 0:
                collect_count = 0
            if forward_count < 0:
                forward_count = 0
            weight_reply = 0.45  # 评论权重
            weight_read = 0.236  # 阅读权重
            weight_like = 0.0  # 点赞权重
            weight_collect = 0.0  # 收藏权重
            weight_forward = 0.314  # 转载权重
            heat = reply_count * weight_reply + read_count * weight_read + like_count * weight_like + collect_count * weight_collect + forward_count * weight_forward
            sql.append(" ('%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s') " %
                       (Tid, event_id, title, add_datetime, publish_datetime, author_id, author_name, content,
                        reply_count, read_count, like_count, collect_count, forward_count, channeltype, channel,
                        correlation, heatwords, heat))
        if sql:
            sql_insert = """INSERT INTO text_select VALUES""" + ','.join(sql) \
                         + """ON DUPLICATE KEY UPDATE reply_count = VALUES(reply_count), read_count = VALUES(read_count), like_count = VALUES(like_count),collect_count = VALUES(collect_count),forward_count = VALUES(forward_count)"""
            session.execute(sql_insert)
            session.commit()
            # sql_insert = """
            #     INSERT INTO text_select (Tid,event_id,title,add_datetime,publish_datetime,author_id,author_name,content,
            #     reply_count,read_count,like_count,collect_count,forward_count,channeltype,channel,correlation,heatwords,heat)
            #     VALUES ('%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s')
            #     ON DUPLICATE KEY UPDATE reply_count = VALUES(reply_count)
            # """ % (Tid, event_id, title, add_datetime, publish_datetime, author_id, author_name, content, reply_count,
            # read_count, like_count, collect_count, forward_count, channeltype, channel, correlation, heatwords, heat)
            # session.execute(sql_insert)
            # session.commit()

    def zhihuSelect(self):
        """
        筛选————知乎
        """
        key = self.key_list
        sql_t = []
        for k in key:
            sql_t.append(" Ttitle LIKE '%s' OR Tcontent LIKE '%s'" % (k, k))
        sql_query = """SELECT * FROM zhihu_text WHERE (""" + 'OR'.join(sql_t) + ") AND Tpublish_datetime > '%s'" % self.start_datetime
        query = session.execute(sql_query)
        sql = []
        # q_news = session.query(NewsText).filter(or_(NewsText.Ttitle.like(key), NewsText.Tcontent.like(key))).all()
        for n in query:
            Tid = n.id + ';' + n.mid
            event_id = self.event_id
            title = n.Ttitle
            add_datetime = n.add_datetime
            publish_datetime = n.Tpublish_datetime
            author_id = n.Tauthor_id
            author_name = n.Tauthor_name
            content = n.Tcontent
            reply_count = n.Treply
            read_count = n.Tread
            like_count = n.Tlike
            collect_count = n.Tfans
            forward_count = 0
            channeltype = 'zhihu'
            channel = 'zhihu.com'
            correlation = ''
            heatwords = ''
            if reply_count < 0:
                reply_count = 0
            if read_count < 0:
                read_count = 0
            if like_count < 0:
                like_count = 0
            if collect_count < 0:
                collect_count = 0
            if forward_count < 0:
                forward_count = 0
            weight_reply = 0.45  # 评论权重
            weight_read = 0.236  # 阅读权重
            weight_like = 0.0  # 点赞权重
            weight_collect = 0.0  # 收藏权重
            weight_forward = 0.314  # 转载权重
            heat = reply_count * weight_reply + read_count * weight_read + like_count * weight_like + collect_count * weight_collect + forward_count * weight_forward
            sql.append(" ('%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s') " %
                       (Tid, event_id, title, add_datetime, publish_datetime, author_id, author_name, content,
                        reply_count, read_count, like_count, collect_count, forward_count, channeltype, channel,
                        correlation, heatwords, heat))
        if sql:
            sql_insert = """INSERT INTO text_select VALUES""" + ','.join(sql) \
                         + """ON DUPLICATE KEY UPDATE reply_count = VALUES(reply_count), read_count = VALUES(read_count), like_count = VALUES(like_count),collect_count = VALUES(collect_count),forward_count = VALUES(forward_count)"""
            session.execute(sql_insert)
            session.commit()
            # sql_insert = """
            #     INSERT INTO text_select (Tid,event_id,title,add_datetime,publish_datetime,author_id,author_name,content,
            #     reply_count,read_count,like_count,collect_count,forward_count,channeltype,channel,correlation,heatwords,heat)
            #     VALUES ('%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s')
            #     ON DUPLICATE KEY UPDATE reply_count = VALUES(reply_count)
            # """ % (Tid, event_id, title, add_datetime, publish_datetime, author_id, author_name, content, reply_count,
            # read_count, like_count, collect_count, forward_count, channeltype, channel, correlation, heatwords, heat)
            # session.execute(sql_insert)
            # session.commit()

    def weiboSelect(self):
        """
        筛选————微博
        """
        key = self.key_list
        sql_t = []
        for k in key:
            sql_t.append(" content LIKE '%s' " % k)
        sql_query = """SELECT * FROM t_sinablog WHERE (""" + 'OR'.join(sql_t) + ") AND publish_datetime > '%s'" % self.start_datetime
        query = session.execute(sql_query)
        sql = []
        # q_news = session.query(NewsText).filter(or_(NewsText.Ttitle.like(key), NewsText.Tcontent.like(key))).all()
        for n in query:
            Tid = n.mid
            event_id = self.event_id
            title = n.content.strip().decode('utf8')[0:19].encode('utf8')
            add_datetime = n.add_datetime
            publish_datetime = n.publish_datetime
            author_id = n.userid
            author_name = n.username
            content = n.content.strip().replace("'","''")
            reply_count = n.comment_count
            read_count = 0
            like_count = n.like_count
            collect_count = 0
            forward_count = n.forward_count
            channeltype = 'weibo'
            channel = 'weibo.com'
            correlation = ''
            heatwords = ''
            if reply_count < 0:
                reply_count = 0
            if read_count < 0:
                read_count = 0
            if like_count < 0:
                like_count = 0
            if collect_count < 0:
                collect_count = 0
            if forward_count < 0:
                forward_count = 0
            weight_reply = 0.45  # 评论权重
            weight_read = 0.236  # 阅读权重
            weight_like = 0.0  # 点赞权重
            weight_collect = 0.0  # 收藏权重
            weight_forward = 0.314  # 转载权重
            heat = reply_count * weight_reply + read_count * weight_read + like_count * weight_like + collect_count * weight_collect + forward_count * weight_forward
            # print title
            # print content
            sql.append(" ('%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s') " %
                       (Tid, event_id, title, add_datetime, publish_datetime, author_id, author_name, content,
                        reply_count, read_count, like_count, collect_count, forward_count, channeltype, channel,
                        correlation, heatwords, heat))
        if sql:
            sql_insert = """INSERT INTO text_select VALUES""" + ','.join(sql) \
                         + """ON DUPLICATE KEY UPDATE reply_count = VALUES(reply_count), read_count = VALUES(read_count), like_count = VALUES(like_count),collect_count = VALUES(collect_count),forward_count = VALUES(forward_count)"""
            session.execute(sql_insert)
            session.commit()

    def weixinSelect(self):
        """
        筛选————微信
        """
        key = self.key_list
        sql_t = []
        for k in key:
            sql_t.append(" title LIKE '%s' OR content LIKE '%s'" % (k, k))
        sql_query = """SELECT * FROM t_weixin WHERE (""" + 'OR'.join(sql_t) + ") AND publish_datetime > '%s'" % self.start_datetime
        query = session.execute(sql_query)
        sql = []
        # q_news = session.query(NewsText).filter(or_(NewsText.Ttitle.like(key), NewsText.Tcontent.like(key))).all()
        for n in query:
            Tid = n.pid + ';' + n.wid
            event_id = self.event_id
            title = n.title
            add_datetime = n.add_datetime
            publish_datetime = n.publish_datetime
            author_id = ''
            author_name = n.author
            content = n.content
            reply_count = 0
            read_count = n.read_count
            like_count = n.like_count
            collect_count = 0
            forward_count = 0
            channeltype = 'weixin'
            channel = 'wx.qq.com'
            correlation = ''
            heatwords = ''
            if reply_count < 0:
                reply_count = 0
            if read_count < 0:
                read_count = 0
            if like_count < 0:
                like_count = 0
            if collect_count < 0:
                collect_count = 0
            if forward_count < 0:
                forward_count = 0
            weight_reply = 0.45  # 评论权重
            weight_read = 0.236  # 阅读权重
            weight_like = 0.0  # 点赞权重
            weight_collect = 0.0  # 收藏权重
            weight_forward = 0.314  # 转载权重
            heat = reply_count * weight_reply + read_count * weight_read + like_count * weight_like + collect_count * weight_collect + forward_count * weight_forward
            sql.append(" ('%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s') " %
                       (Tid, event_id, title, add_datetime, publish_datetime, author_id, author_name, content,
                        reply_count, read_count, like_count, collect_count, forward_count, channeltype, channel,
                        correlation, heatwords, heat))
        if sql:
            sql_insert = """INSERT INTO text_select VALUES""" + ','.join(sql) \
                         + """ON DUPLICATE KEY UPDATE reply_count = VALUES(reply_count), read_count = VALUES(read_count), like_count = VALUES(like_count),collect_count = VALUES(collect_count),forward_count = VALUES(forward_count)"""
            session.execute(sql_insert)
            session.commit()
            # sql_insert = """
            #     INSERT INTO text_select (Tid,event_id,title,add_datetime,publish_datetime,author_id,author_name,content,
            #     reply_count,read_count,like_count,collect_count,forward_count,channeltype,channel,correlation,heatwords,heat)
            #     VALUES ('%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s')
            #     ON DUPLICATE KEY UPDATE reply_count = VALUES(reply_count)
            # """ % (Tid, event_id, title, add_datetime, publish_datetime, author_id, author_name, content, reply_count,
            # read_count, like_count, collect_count, forward_count, channeltype, channel, correlation, heatwords, heat)
            # session.execute(sql_insert)
            # session.commit()

class Growth():
    """
    文章指标增长模块
    """
    def __init__(self):
        self.t = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(time.time()))

    def updateGrowth(self):
        """将text_select文章按每小时更新到text_growth中"""
        CrawlTime = self.t
        sql_query = """
                    SELECT Tid,event_id,reply_count,read_count,like_count,collect_count,forward_count,channeltype,channel,heat
                    FROM text_select
                """
        query = session.execute(sql_query)
        sql = []
        for q in query:
            sql.append(" ('%s','%s','%s','%s','%s','%s','%s','%s','%s','%s','%s') " % (q[0],q[1],CrawlTime,q[2],q[3],q[4],q[5],q[6],q[7],q[8],q[9]))
        sql_insert = """INSERT INTO text_growth VALUES""" + ','.join(sql)
        session.execute(sql_insert)
        session.commit()

    def updateArticleHeat(self):
        """将文章热度计算并更新到text_select和text_growth中"""
        weight_reply = 0.45  # 评论权重
        weight_read = 0.236  # 阅读权重
        weight_like = 0.0  # 点赞权重
        weight_collect = 0.0  # 收藏权重
        weight_forward = 0.314  # 转载权重
        hour = 4  # 多少小时内
        times = 4  # 取多少次数
        sql_query1  = " SELECT Tid, event_id, crawl_datetime, channel, a4 / times AS heat FROM "
        sql_query1 += " (SELECT Tid, event_id, crawl_datetime, channel, SUM(a1+ a2+ a3) AS a4, COUNT(*) as times FROM "
        sql_query1 += " (SELECT Tid, event_id, crawl_datetime, channel, reply_count * %f AS a1, read_count * %f AS a2, forward_count * %f AS a3 " % (weight_reply, weight_read, weight_forward)
        sql_query1 += " FROM text_growth "
        sql_query1 += " WHERE crawl_datetime > DATE_SUB(NOW(),INTERVAL %d HOUR) " % hour
        sql_query1 += " ORDER BY crawl_datetime DESC) t1 "
        sql_query1 += " GROUP BY Tid, event_id, channel) t2 "
        sql_query1 += " WHERE times = %d " % times
        query1 = session.execute(sql_query1)
        values1 = []
        values2 = []
        if not query1:
            for q1 in query1:
                value1 = " ('%s','%s','%s','%s',%.2f) " % (q1[0],q1[1],q1[2],q1[3],q1[4])
                value2 = " ('%s','%s','%s',%.2f) " % (q1[0], q1[1], q1[3], q1[4])
                values1.append(value1)
                values2.append(value2)
            sql_insert1  = " INSERT INTO text_growth (Tid, event_id, crawl_datetime, channel, heat) VALUES "
            sql_insert1 += ','.join(values1)
            sql_insert1 += " ON DUPLICATE KEY UPDATE heat=VALUES(heat)"
            sql_insert2 = " INSERT INTO text_select (Tid, event_id, channel, heat) VALUES "
            sql_insert2 += ','.join(values2)
            sql_insert2 += " ON DUPLICATE KEY UPDATE heat=VALUES(heat)"
            session.execute(sql_insert1)  # 更新到text_growth
            session.execute(sql_insert2)  # 更新到text_select
            session.commit()

class Indication():
    """
    事件指标计算（热度）
    """
    def __init__(self):
        self.t = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(time.time()))

    def insertEventHeat(self):
        weight_reply = 0.45  # 评论权重
        weight_read = 0.236  # 阅读权重
        weight_like = 0.0  # 点赞权重
        weight_collect = 0.0  # 收藏权重
        weight_forward = 0.314  # 转载权重
        sql_insert = """
            insert ignore into event_statistics (event_id,crawl_datetime,reply_total,read_total,like_total,collect_total,original_total,forward_total,event_heat)
            select event_id,crawl_datetime,sum(reply_count) as reply_total,sum(read_count) as read_total,
            sum(like_count) as like_total,sum(collect_count) as collect_total,count(*) as original_total,
            sum(forward_count) as forward_total,sum(reply_count * %f + read_count * %f + forward_count * %f ) as event_heat
            from text_growth
            group by crawl_datetime, event_id
            order by crawl_datetime desc
        """ % (weight_reply,weight_read,weight_forward)
        session.execute(sql_insert)
        session.commit()

    def updateEventHeat(self):
        weight_reply = 0.45  # 评论权重
        weight_read = 0.236  # 阅读权重
        weight_like = 0.0  # 点赞权重
        weight_collect = 0.0  # 收藏权重
        weight_forward = 0.314  # 转载权重
        hour = 4  # 多少小时内
        times = 4  # 取多少次数
        sql_query1  = " INSERT INTO event_statistics (event_id, crawl_datetime, event_heat) "
        sql_query1 += " SELECT event_id, crawl_datetime, a4 / times AS event_heat FROM "
        sql_query1 += " (SELECT event_id, crawl_datetime, SUM(a1+ a2+ a3) AS a4, COUNT(*) as times FROM "
        sql_query1 += " (SELECT event_id, crawl_datetime, reply_total * %f AS a1, read_total * %f AS a2, forward_total * %f AS a3 " % (weight_reply, weight_read, weight_forward)
        sql_query1 += " FROM event_statistics "
        sql_query1 += " WHERE crawl_datetime > DATE_SUB(NOW(),INTERVAL %d HOUR) " % hour
        sql_query1 += " ORDER BY crawl_datetime DESC) t1 "
        sql_query1 += " GROUP BY event_id) t2 "
        sql_query1 += " WHERE times = %d " % times
        sql_query1 += " ON DUPLICATE KEY UPDATE event_heat=VALUES(event_heat) "
        session.execute(sql_query1)
        session.commit()

    def insertHeatWords(self):
        sql_query1 = " SELECT event_id FROM event_list "
        query1 = session.execute(sql_query1)
        sql = []
        for q in query1:
            sql_query2 = " SELECT content FROM text_select WHERE event_id=%d " % q.event_id
            query2 = session.execute(sql_query2)
            stopWords = ['...', 'via', 'video', u'中山大学', u'中大', u'视频', u'校园', u'学校', u'高校', u'大学', u'学生', u'师生', u'广州',
                         u'深圳', u'如何', u'评价', u'链接', u'网页', u'原标题']
            # 将所有文章内容连接
            longcontent = '\n'.join(map(lambda x: x[0], query2))
            for w in stopWords:
                longcontent = longcontent.replace(w, '')
            # 分析关键词
            import jieba.analyse
            tags = jieba.analyse.extract_tags(longcontent, topK=15, withWeight=True)
            for t in tags:
                sql.append(" ('%s','%s','%s','%s') "% (q.event_id, self.t ,t[0] ,t[1]))
                # print q.event_id, self.t, t[0], t[1]

        # 查询所有
        sql_query3 = " SELECT content FROM text_select "
        query3 = session.execute(sql_query3)
        stopWords = ['...', 'via', 'video', u'中山大学', u'中大', u'视频', u'校园', u'学校', u'高校', u'大学', u'学生', u'师生', u'广州',
                     u'深圳', u'如何', u'评价', u'链接', u'网页', u'原标题']
        longcontent = '\n'.join(map(lambda x: x[0], query3))
        for w in stopWords:
            longcontent = longcontent.replace(w, '')
        alls = jieba.analyse.extract_tags(longcontent, topK=20, withWeight=True)
        for a in alls:
            sql.append(" ('%s','%s','%s','%s') " % (0, self.t, a[0], a[1]))

        sql_insert = " INSERT INTO event_heatwords VALUES " + ','.join(sql)
        session.execute(sql_insert)
        session.commit()
        # print ','.join(sql)

class Report():
    """
    报告生成与下载
    """
    def __init__(self):
        self.t = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(time.time()))
        self.d = time.strftime('%Y%m%d', time.localtime(time.time()))

    def overviewReport(self):
        from docx.shared import Pt
        from docx.oxml.ns import qn
        from docx.enum.style import WD_STYLE_TYPE
        print "start generating overview report"
        document = Document()
        songti = u'宋体'
        h1 = document.add_heading(u'第一章 事件统计数据一览表', level=1)


        paragraph = document.add_paragraph()
        run = paragraph.add_run(u'中文内容')
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.name = songti
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), songti)
        document.add_page_break()
        path = r'C:\Users\LB\Desktop\%s%s.docx' % (self.d, u'报告(所有事件)')
        document.save(path)


    def singleReport(self):
        print "start generating single report"

if __name__ == '__main__':
    print "——————————start——————————"
    # Growth().updateGrowth()
    # Growth().updateArticleHeat()
    # Indication().insertEventHeat()
    # Indication().updateEventHeat()
    # Indication().insertHeatWords()
    # Report().overviewReport()
    # Processing(5,'2017-06-22 16:18:54','2017-06-20 00:00:00','中央巡视组','').weiboSelect()