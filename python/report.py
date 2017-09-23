# -*- coding:utf-8 -*-
import matplotlib
matplotlib.use('Agg')
import os
import time
from pylab import *
from db_init import *
from gevent import monkey, pool
monkey.patch_all()
pool = pool.Pool(10)

mpl.rcParams['font.sans-serif'] = ['Microsoft YaHei']  # 指定默认字体
mpl.rcParams['axes.unicode_minus'] = False  # 解决保存图像是负号'-'显示为方块的问题

class allEventReport():
    """
    报告————事件总览
    """
    def __init__(self):
        self.root_path = os.getcwd()  # 根目录
        self.image_path = self.root_path + r'/image'  # image目录
        self.report_path = self.root_path + r'/reports'  # reports目录
        self.t = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(time.time()))
        self.d = time.strftime('%Y%m%d', time.localtime(time.time()))

    def Controller(self):
        sql = "SELECT * FROM event_list"
        resultList = session.execute(sql)
        for r in resultList:
            event_id = r['event_id']
            pool.spawn(self.eventHeatChart, event_id)
            pool.spawn(self.mediaReplyChart, event_id)
        pool.join()
        self.newDocx()  # 适用python-docx

    def eventHeatChart(self, event_id):
        """
        事件热度变化值————折线图
        :return:
        """
        sql = """
            SELECT a.event_id, b.event_name, DATE_FORMAT(crawl_datetime,'%%Y/%%m/%%d') AS crawl_date, event_heat FROM
            (SELECT event_id, crawl_datetime, event_heat FROM event_statistics_copy WHERE event_id = %d ORDER BY crawl_datetime DESC) a
            LEFT JOIN event_list b
            ON a.event_id = b.event_id
            GROUP BY TO_DAYS(crawl_datetime)
        """ % event_id
        resultList = session.execute(sql)
        xList = []
        yList = []
        heatDetaList = []  # 热度变化值列表
        event_name = ''
        for r in resultList:
            event_id = r[0]
            event_name = r[1]
            xList.append(r[2])
            yList.append(r[3])
        if event_name != '':
            it = iter(yList)
            n1 = float(it.next())
            for n2 in it:
                if n1 != 0:
                    heatDetaList.append((n2 - n1) / n1)
                else:
                    heatDetaList.append(n2)
                n1 = float(n2)

            fig = plt.figure(figsize=(8, 4))  # 生成折线图
            ax = fig.add_subplot(111)  # 建立axis对象在第1行第1列编号为1
            X1 = range(0, len(heatDetaList))  # x轴数据
            Y1 = heatDetaList  # y轴数据
            ax.plot(X1, Y1)
            ax.spines['top'].set_color('none')  # 上端边框颜色为透明
            ax.spines['right'].set_color('none')  # 右端边框颜色为透明
            ax.yaxis.grid(True, which='major', linestyle="--")  # y坐标轴的网格使用主刻度
            ax.set_xticks(range(0, len(heatDetaList)))  # x轴坐标设置
            ax.set_xticklabels(xList[1:], rotation=-30)  # x轴坐标标签设置
            ax.set_xlabel(u'日期')  # x轴标签
            ax.set_ylabel(u'热度')  # y轴标签
            ax.set_title(event_name + u'——事件热度')  # 标题设置
            plt.tight_layout()  # 画布布局自适应
            # plt.show()
            # fig.show()  # 生成图表
            savePath = self.image_path + r'/detaHeat_%s.jpg' % event_id
            fig.savefig(savePath)  # 保存图表

    def mediaReplyChart(self, event_id):
        """
        媒体评论数————饼图
        :return:
        """
        sql = """
            SELECT a.event_id, b.event_name, chi AS media, SUM(reply_count) AS reply_count
            FROM text_select_copy a
            LEFT JOIN event_list b
            ON a.event_id = b.event_id
            LEFT JOIN translation
            ON channel = eng
            WHERE a.event_id = %d
            GROUP BY channel
            ORDER BY reply_count DESC
        """ % event_id
        resultList = session.execute(sql)
        xList = []
        yList = []
        event_name = ''
        for r in resultList:
            event_id = r[0]
            event_name = r[1]
            xList.append(r[2])
            yList.append(r[3])
        if event_name != '':
            xList1 = []
            for i in range(0, len(yList)):
                if sum(yList) != 0:
                    pct = '%.2f%%' % ((yList[i] / sum(yList)) * 100)
                else:
                    pct = 0.0
                xList1.append(xList[i] + ' ' + str(pct))
            fig = plt.figure(figsize=(6, 4))
            ax = fig.add_subplot(111)  # 建立axis对象在第1行第1列编号为1
            labels = xList1
            sizes = yList
            ax.pie(sizes, labeldistance=1, radius=0.9, shadow=False, startangle=90)
            ax.set(aspect='equal')
            ax.set(title=event_name + u'——评论')
            plt.legend(labels, bbox_to_anchor=(0.9, 0.9), loc='best', borderaxespad=0)
            plt.tight_layout()
            # plt.show()
            # fig.show()
            savePath = self.image_path + r'/mediaReply_%s.jpg' % event_id
            fig.savefig(savePath)

    def newDocx(self):
        from docx import Document
        from docx.shared import Pt
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        document = Document(self.report_path + r'/all_demo1.docx')

        # 第一章
        para1 = document.paragraphs[0]
        para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = para1.add_run(u'第一章 事件统计数据一览表')
        run1.font.size = Pt(16)
        run1.bold = True
        table1 = document.add_table(rows=1, cols=12, style='Table Grid')
        table1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells = table1.rows[0].cells
        hdr_cells[0].width = Inches(1.8)
        hdr_cells[1].width = Inches(1.1)
        hdr_cells[2].width = Inches(1.1)
        hdr_cells[3].width = Inches(1.1)
        hdr_cells[4].width = Inches(1.1)
        hdr_cells[5].width = Inches(0.8)
        hdr_cells[6].width = Inches(0.6)
        hdr_cells[7].width = Inches(0.6)
        hdr_cells[8].width = Inches(0.6)
        hdr_cells[9].width = Inches(0.6)
        hdr_cells[10].width = Inches(0.6)
        hdr_cells[11].width = Inches(0.6)
        hdr_cells[0].paragraphs[0].add_run(u'事件名').bold = True
        hdr_cells[1].paragraphs[0].add_run(u'添加时间').bold = True
        hdr_cells[2].paragraphs[0].add_run(u'监测开始时间').bold = True
        hdr_cells[3].paragraphs[0].add_run(u'监测结束时间').bold = True
        hdr_cells[4].paragraphs[0].add_run(u'更新时间').bold = True
        hdr_cells[5].paragraphs[0].add_run(u'事件热度').bold = True
        hdr_cells[6].paragraphs[0].add_run(u'总评论').bold = True
        hdr_cells[7].paragraphs[0].add_run(u'总阅读').bold = True
        hdr_cells[8].paragraphs[0].add_run(u'总点赞').bold = True
        hdr_cells[9].paragraphs[0].add_run(u'总收藏').bold = True
        hdr_cells[10].paragraphs[0].add_run(u'总原创').bold = True
        hdr_cells[11].paragraphs[0].add_run(u'总转发').bold = True
        sql_1 = """
                    SELECT event_name,add_datetime,start_datetime,crawl_datetime,event_heat,reply_total,read_total,like_total,collect_total,original_total,forward_total
                    FROM event_list a INNER JOIN (SELECT * FROM event_statistics_copy ORDER BY crawl_datetime DESC) b ON a.event_id=b.event_id
                    GROUP BY a.event_id ORDER BY b.event_heat DESC
                """
        resultList1 = session.execute(sql_1)
        total_reply, total_read, total_like, total_collect, total_original, total_forward = 0, 0, 0, 0, 0, 0
        for r in resultList1:
            cells = table1.add_row().cells
            cells[0].width = Inches(1.8)
            cells[1].width = Inches(1.1)
            cells[2].width = Inches(1.1)
            cells[3].width = Inches(1.1)
            cells[4].width = Inches(1.1)
            cells[5].width = Inches(0.8)
            cells[6].width = Inches(0.6)
            cells[7].width = Inches(0.6)
            cells[8].width = Inches(0.6)
            cells[9].width = Inches(0.6)
            cells[10].width = Inches(0.6)
            cells[11].width = Inches(0.6)
            cells[0].text = r[0]
            cells[1].text = str(r[1])
            cells[2].text = str(r[2])
            cells[3].text = ''
            cells[4].text = str(r[3])
            cells[5].text = str(r[4])
            cells[6].text = str(r[5])
            cells[7].text = str(r[6])
            cells[8].text = str(r[7])
            cells[9].text = str(r[8])
            cells[10].text = str(r[9])
            cells[11].text = str(r[10])
            total_reply += r[5]
            total_read += r[6]
            total_like += r[7]
            total_collect += r[8]
            total_original += r[9]
            total_forward += r[10]
        cells = table1.add_row().cells
        cells[0].width = Inches(1.8)
        cells[1].width = Inches(1.1)
        cells[2].width = Inches(1.1)
        cells[3].width = Inches(1.1)
        cells[4].width = Inches(1.1)
        cells[5].width = Inches(0.8)
        cells[6].width = Inches(0.6)
        cells[7].width = Inches(0.6)
        cells[8].width = Inches(0.6)
        cells[9].width = Inches(0.6)
        cells[10].width = Inches(0.6)
        cells[11].width = Inches(0.6)
        cells[0].merge(cells[5])
        cells[0].paragraphs[0].add_run(u'合计').bold = True
        cells[6].text = str(total_reply)
        cells[7].text = str(total_read)
        cells[8].text = str(total_like)
        cells[9].text = str(total_collect)
        cells[10].text = str(total_original)
        cells[11].text = str(total_forward)
        document.add_page_break()

        # 第二章
        para2 = document.add_paragraph()
        para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = para2.add_run(u'第二章 最近一周热词分析表')
        run2.font.size = Pt(16)
        run2.bold = True
        table2 = document.add_table(rows=1, cols=2, style='Table Grid')
        table2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells = table2.rows[0].cells
        hdr_cells[0].width = Inches(1.0)
        hdr_cells[1].width = Inches(1.0)
        hdr_cells[0].paragraphs[0].add_run(u'热词').bold = True
        hdr_cells[1].paragraphs[0].add_run(u'热度').bold = True
        sql_2 = """
                    SELECT crawl_datetime, words_name, words_weight
                    FROM event_heatwords
                    WHERE event_id = 0 AND crawl_datetime = (SELECT MAX(crawl_datetime) FROM event_heatwords WHERE event_id = 0)
                    ORDER BY words_weight DESC
                """
        resultList2 = session.execute(sql_2)
        for r in resultList2:
            cells = table2.add_row().cells
            cells[0].width = Inches(1.0)
            cells[1].width = Inches(1.0)
            cells[0].text = r[1]
            cells[1].text = str(r[2])
        document.add_page_break()

        # 第三章
        para3 = document.add_paragraph()
        para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = para3.add_run(u'第三章 事件热度变化曲线图')
        run3.font.size = Pt(16)
        run3.bold = True
        items = os.listdir(self.image_path)  # image目录文件
        for names in items[::-1]:
            if 'detaHeat_' in names:
                detaHeat_path = self.image_path + "//" + names
                document.add_picture(detaHeat_path, width=Inches(9))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_page_break()

        # 第四章
        para4 = document.add_paragraph()
        para4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run4 = para4.add_run(u'第四章 事件评论数媒体分布图')
        run4.font.size = Pt(16)
        run4.bold = True
        items = os.listdir(self.image_path)  # image目录文件
        for names in items[::-1]:
            if 'mediaReply_' in names:
                mediaReply_path = self.image_path + "//" + names
                pic4 = document.add_picture(mediaReply_path, width=Inches(8))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        report_name = u'事件总览报告' + self.d + '.docx'
        document.save(self.report_path + '//' + report_name)

class onlyEventReport():
    """
    报告————单事件
    """
    def __init__(self):
        self.root_path = os.getcwd()  # 根目录
        self.image_path = self.root_path + r'/image'  # image目录
        self.report_path = self.root_path + r'/reports'  # reports目录
        self.t = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(time.time()))
        self.d = time.strftime('%Y%m%d', time.localtime(time.time()))

    def Controller(self, event_id):
        self.transmissionChart(event_id)
        self.readChart(event_id)
        self.eventHeatChart(event_id)
        self.ori_forwardChart(event_id)
        self.provinceChart(event_id)
        self.cityChart(event_id)
        self.newDocx(event_id)  # python-docx

    def transmissionChart(self, event_id):
        """
        传播链————折线图
        :return:
        """
        sql = """
            SELECT chi, publish_datetime
            FROM text_select_copy
            LEFT JOIN translation
            ON channel = eng
            WHERE event_id = %d
            GROUP BY channel
            ORDER BY publish_datetime
        """ % event_id
        resultList = session.execute(sql)
        xList = []
        yList = []
        for r in resultList:
            channel = r[0]
            publish_datetime = r[1]
            xList.append(channel)
            yList.append(publish_datetime)

        fig = plt.figure(figsize=(8, 4))  # 生成折线图
        ax = fig.add_subplot(111)  # 建立axis对象在第1行第1列编号为1
        X1 = range(0, len(xList))  # x轴数据
        Y1 = yList  # y轴数据
        ax.plot(X1, Y1, '.-')
        ax.spines['top'].set_color('none')  # 上端边框颜色为透明
        ax.spines['right'].set_color('none')  # 右端边框颜色为透明
        ax.yaxis.grid(True, which='major', linestyle="--")  # y坐标轴的网格使用主刻度
        ax.set_xticks(range(0, len(xList)))  # x轴坐标设置
        ax.set_xticklabels(xList, rotation=-30)  # x轴坐标标签设置
        # ax.set_yticks(range(0, len(yList)))  # x轴坐标设置
        # ax.set_yticklabels(yList, rotation=-30)  # x轴坐标标签设置
        ax.set_xlabel(u'媒体')  # x轴标签
        ax.set_ylabel(u'时间')  # y轴标签
        ax.set_title(u'传播链')  # 标题设置
        plt.tight_layout()  # 画布布局自适应
        # plt.show()
        # fig.show()  # 生成图表
        savePath = self.image_path + r'/transmission_%s.jpg' % event_id
        fig.savefig(savePath)  # 保存图表

    def readChart(self, event_id):
        """
        阅读量————折线图
        :return:
        """
        sql = """
            SELECT DATE_FORMAT(crawl_datetime,'%%Y/%%m/%%d'), read_total
            FROM (SELECT * FROM event_statistics_copy ORDER BY crawl_datetime) a
            WHERE event_id = %d
            GROUP BY TO_DAYS(crawl_datetime)
        """ % event_id
        resultList = session.execute(sql)
        xList = []
        yList = []
        for r in resultList:
            crawl_datetime = r[0]
            read_total = r[1]
            xList.append(crawl_datetime)
            yList.append(read_total)

        fig = plt.figure(figsize=(8, 4))  # 生成折线图
        ax = fig.add_subplot(111)  # 建立axis对象在第1行第1列编号为1
        X1 = range(0, len(xList))  # x轴数据
        Y1 = yList  # y轴数据
        ax.plot(X1, Y1, '-')
        ax.spines['top'].set_color('none')  # 上端边框颜色为透明
        ax.spines['right'].set_color('none')  # 右端边框颜色为透明
        ax.yaxis.grid(True, which='major', linestyle="--")  # y坐标轴的网格使用主刻度
        ax.set_xticks(range(0, len(xList)))  # x轴坐标设置
        ax.set_xticklabels(xList, rotation=-30)  # x轴坐标标签设置
        ax.set_xlabel(u'日期')  # x轴标签
        ax.set_ylabel(u'阅读量')  # y轴标签
        ax.set_title(u'阅读总量')  # 标题设置
        plt.tight_layout()  # 画布布局自适应
        # plt.show()
        # fig.show()  # 生成图表
        savePath = self.image_path + r'/read_%s.jpg' % event_id
        fig.savefig(savePath)  # 保存图表

    def eventHeatChart(self, event_id):
        """
        事件热度变化值————折线图
        :return:
        """
        sql = """
            SELECT a.event_id, b.event_name, DATE_FORMAT(crawl_datetime,'%%Y/%%m/%%d') AS crawl_date, event_heat FROM
            (SELECT event_id, crawl_datetime, event_heat FROM event_statistics_copy WHERE event_id = %d ORDER BY crawl_datetime DESC) a
            LEFT JOIN event_list b
            ON a.event_id = b.event_id
            GROUP BY TO_DAYS(crawl_datetime)
        """ % event_id
        resultList = session.execute(sql)
        xList = []
        yList = []
        heatDetaList = []  # 热度变化值列表
        event_name = ''
        for r in resultList:
            event_id = r[0]
            event_name = r[1]
            xList.append(r[2])
            yList.append(r[3])
        if event_name != '':
            it = iter(yList)
            n1 = float(it.next())
            for n2 in it:
                heatDetaList.append((n2 - n1) / n1)
                n1 = float(n2)

            fig = plt.figure(figsize=(8, 4))  # 生成折线图
            ax = fig.add_subplot(111)  # 建立axis对象在第1行第1列编号为1
            X1 = range(0, len(heatDetaList))  # x轴数据
            Y1 = heatDetaList  # y轴数据
            ax.plot(X1, Y1)
            ax.spines['top'].set_color('none')  # 上端边框颜色为透明
            ax.spines['right'].set_color('none')  # 右端边框颜色为透明
            ax.yaxis.grid(True, which='major', linestyle="--")  # y坐标轴的网格使用主刻度
            ax.set_xticks(range(0, len(heatDetaList)))  # x轴坐标设置
            ax.set_xticklabels(xList[1:], rotation=-30)  # x轴坐标标签设置
            ax.set_xlabel(u'日期')  # x轴标签
            ax.set_ylabel(u'热度')  # y轴标签
            ax.set_title(event_name + u'——事件热度')  # 标题设置
            plt.tight_layout()  # 画布布局自适应
            # plt.show()
            # fig.show()  # 生成图表
            savePath = self.image_path + r'/detaHeat_%s.jpg' % event_id
            fig.savefig(savePath)  # 保存图表

    def ori_forwardChart(self, event_id):
        """
        原创转发比————饼图
        :return:
        """
        sql = """
            SELECT original_total, forward_total
            FROM event_statistics_copy
            WHERE event_id = %d
            ORDER BY crawl_datetime DESC
            LIMIT 1
        """ % event_id
        resultList = session.execute(sql)
        yList = []
        for r in resultList:
            original_total = r[0]
            forward_total = r[1]
            yList.append(original_total)
            yList.append(forward_total)
        fig = plt.figure(figsize=(6, 4))
        ax = fig.add_subplot(111)  # 建立axis对象在第1行第1列编号为1
        sizes = yList
        ax.pie(sizes, labeldistance=1, radius=0.9, shadow=False, startangle=90, autopct = '%.1f%%', pctdistance = 1.1)
        ax.set(aspect='equal')
        ax.set(title=u'原创转发比')
        plt.legend([u'原创',u'转发'], bbox_to_anchor=(1.2, 0.9), loc='best', borderaxespad=0)
        plt.tight_layout()
        # plt.show()
        # fig.show()
        savePath = self.image_path + r'/ori_forward_%s.jpg' % event_id
        fig.savefig(savePath)

    def provinceChart(self, event_id):
        """
        网民地理（全国）————柱状图
        :return:
        """
        sql = u"""
            SELECT region, count(*) AS number
            FROM (
            SELECT substring_index(substring_index(ip_address,';',2),';',-1) AS region
            FROM text_select_comment
            WHERE ip_address LIKE '中国%%' AND event_id = %d
            ) AS a
            GROUP BY region
            ORDER BY number DESC
            LIMIT 10
        """ % event_id
        resultList = session.execute(sql)
        xList = []
        yList = []
        for r in resultList:
            region = r[0]
            number = r[1]
            xList.append(region)
            yList.append(number)
        fig = plt.figure(figsize=(8, 4))  # 生成折线图
        ax = fig.add_subplot(111)  # 建立axis对象在第1行第1列编号为1
        X1 = range(0, len(xList))  # x轴数据
        Y1 = yList  # y轴数据
        ax.bar(X1, Y1)
        ax.spines['top'].set_color('none')  # 上端边框颜色为透明
        ax.spines['right'].set_color('none')  # 右端边框颜色为透明
        ax.yaxis.grid(True, which='major', linestyle="--")  # y坐标轴的网格使用主刻度
        ax.set_xticks(range(0, len(xList)))  # x轴坐标设置
        ax.set_xticklabels(xList, rotation=-30)  # x轴坐标标签设置
        ax.set_xlabel(u'省份')  # x轴标签
        ax.set_ylabel(u'人数')  # y轴标签
        ax.set_title(u'地理分布（全国）')  # 标题设置
        for a, b in zip(X1, Y1):
            plt.text(a, b + 0.1, '%.0f' % b, ha='center', va='bottom')
        plt.tight_layout()  # 画布布局自适应
        # plt.show()
        # fig.show()  # 生成图表
        savePath = self.image_path + r'/province_%s.jpg' % event_id
        fig.savefig(savePath)  # 保存图表

    def cityChart(self, event_id):
        """
        网民地理（省内）————柱状图
        :return:
        """
        sql = u"""
            SELECT city, count(*) AS number
            FROM (
            SELECT substring_index(ip_address,';',-1) AS city
            FROM text_select_comment
            WHERE ip_address LIKE '中国;广东;%%' AND event_id = %d
            ) AS a
            WHERE city <> ''
            GROUP BY city
            ORDER BY number DESC
            LIMIT 10
        """ % event_id
        resultList = session.execute(sql)
        xList = []
        yList = []
        for r in resultList:
            region = r[0]
            number = r[1]
            xList.append(region)
            yList.append(number)
        fig = plt.figure(figsize=(8, 4))  # 生成折线图
        ax = fig.add_subplot(111)  # 建立axis对象在第1行第1列编号为1
        X1 = range(0, len(xList))  # x轴数据
        Y1 = yList  # y轴数据
        ax.bar(X1, Y1)
        ax.spines['top'].set_color('none')  # 上端边框颜色为透明
        ax.spines['right'].set_color('none')  # 右端边框颜色为透明
        ax.yaxis.grid(True, which='major', linestyle="--")  # y坐标轴的网格使用主刻度
        ax.set_xticks(range(0, len(xList)))  # x轴坐标设置
        ax.set_xticklabels(xList, rotation=-30)  # x轴坐标标签设置
        ax.set_xlabel(u'城市')  # x轴标签
        ax.set_ylabel(u'人数')  # y轴标签
        ax.set_title(u'地理分布（省内）')  # 标题设置
        for a, b in zip(X1, Y1):
            plt.text(a, b + 0.1, '%.0f' % b, ha='center', va='bottom')
        plt.tight_layout()  # 画布布局自适应
        # plt.show()
        # fig.show()  # 生成图表
        savePath = self.image_path + r'/city_%s.jpg' % event_id
        fig.savefig(savePath)  # 保存图表

    def newDocx(self, event_id):
        from docx import Document
        from docx.shared import Pt
        from docx.shared import Inches
        from docx.oxml.ns import qn
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        items = os.listdir(self.image_path)  # image目录文件
        document = Document(self.report_path + r'/event_demo2.docx')
        # 第一章
        para1 = document.paragraphs[0]
        para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = para1.add_run(u'第一章 事件属性')
        run1.font.size = Pt(16)
        run1.bold = True
        table1 = document.add_table(rows=4, cols=2, style='Table Grid')
        hdr_cells = table1.columns[0].cells
        hdr_cells[0].width = Inches(1.3)
        hdr_cells[1].width = Inches(1.3)
        hdr_cells[2].width = Inches(1.3)
        hdr_cells[3].width = Inches(1.3)
        hdr_cells[0].paragraphs[0].add_run(u'事件名').bold = True
        hdr_cells[1].paragraphs[0].add_run(u'添加时间').bold = True
        hdr_cells[2].paragraphs[0].add_run(u'监测开始时间').bold = True
        hdr_cells[3].paragraphs[0].add_run(u'监测结束时间').bold = True
        t1_sql = """
                    SELECT event_name,add_datetime,start_datetime
                    FROM event_list
                    WHERE event_id = %d
                """ % event_id
        t1_resultList = session.execute(t1_sql)
        for r in t1_resultList:
            event_name = r[0]
            cells = table1.columns[1].cells
            cells[0].width = Inches(6)
            cells[1].width = Inches(6)
            cells[2].width = Inches(6)
            cells[3].width = Inches(6)
            cells[0].text = r[0]
            cells[1].text = str(r[1])
            cells[2].text = str(r[2])
            cells[3].text = ''
            break

        # 第二章
        para2 = document.add_paragraph()
        para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = para2.add_run(u'第二章 事件传播链')
        run2.font.size = Pt(16)
        run2.bold = True
        table2 = document.add_table(rows=1, cols=4, style='Table Grid')
        hdr_cells = table2.rows[0].cells
        hdr_cells[0].width = Inches(1.5)
        hdr_cells[1].width = Inches(1.0)
        hdr_cells[2].width = Inches(4.0)
        hdr_cells[3].width = Inches(0.8)
        hdr_cells[0].paragraphs[0].add_run(u'各媒体最早出现时间').bold = True
        hdr_cells[1].paragraphs[0].add_run(u'媒体名称').bold = True
        hdr_cells[2].paragraphs[0].add_run(u'标题').bold = True
        hdr_cells[3].paragraphs[0].add_run(u'文章热度').bold = True
        t2_sql = """
                    SELECT publish_datetime as start_datetime, chi, title, heat
                    FROM (SELECT * FROM text_select_copy ORDER BY publish_datetime) a
                    LEFT JOIN (SELECT * FROM text_analysis ORDER BY crawl_datetime DESC) b
                    ON a.Tid = b.Tid AND a.event_id = b.event_id AND a.channel = b.channel
                    LEFT JOIN translation
                    ON a.channel = eng
                    WHERE a.event_id = %d
                    GROUP BY a.channel
                    ORDER BY start_datetime
                """ % event_id
        t2_resultList = session.execute(t2_sql)
        for r in t2_resultList:
            cells = table2.add_row().cells
            cells[0].width = Inches(1.5)
            cells[1].width = Inches(1.0)
            cells[2].width = Inches(4.0)
            cells[3].width = Inches(0.8)
            cells[0].text = str(r[0])
            cells[1].text = r[1]
            cells[2].text = r[2]
            cells[3].text = str(r[3])
        table2.autofit = False
        table2.columns[1].width = Inches(1.0)
        for names in items[::-1]:
            if 'transmission_' + str(event_id) + '.jpg' == names:
                transmission_path = self.image_path + "//" + names
                document.add_picture(transmission_path, width=Inches(7))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_page_break()

        # 第三章
        para3 = document.add_paragraph()
        para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = para3.add_run(u'第三章 事件传播波形图')
        run3.font.size = Pt(16)
        run3.bold = True
        for names in items[::-1]:
            if 'read_' + str(event_id) + '.jpg' == names:
                read_path = self.image_path + "//" + names
                document.add_picture(read_path, width=Inches(7))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para3_1 = document.add_paragraph()
        para3_1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3_1 = para3_1.add_run(u'具体数据见：《单个事件各文章报告》')
        run3_1.font.size = Pt(10.5)
        run3_1.font.name = u'黑体'
        r3_1 = run3_1._element
        r3_1.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
        document.add_page_break()

        # 第四章
        para4 = document.add_paragraph()
        para4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run4 = para4.add_run(u'第四章 事件热度变化图')
        run4.font.size = Pt(16)
        run4.bold = True
        for names in items[::-1]:
            if 'detaHeat_' + str(event_id) + '.jpg' == names:
                detaHeat_path = self.image_path + "//" + names
                document.add_picture(detaHeat_path, width=Inches(7))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para4_1 = document.add_paragraph()
        para4_1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run4_1 = para4_1.add_run(u'具体数据见：《单个事件各文章报告》')
        run4_1.font.size = Pt(10.5)
        run4_1.font.name = u'黑体'
        r4_1 = run4_1._element
        r4_1.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
        document.add_page_break()

        # 第五章
        para5 = document.add_paragraph()
        para5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run5 = para5.add_run(u'第五章 事件原创和转发比')
        run5.font.size = Pt(16)
        run5.bold = True
        table5 = document.add_table(rows=2, cols=7, style='Table Grid')
        hdr_cells = table5.rows[0].cells
        hdr_cells[0].paragraphs[0].add_run(u'总原创数').bold = True
        hdr_cells[1].paragraphs[0].add_run(u'总转发数').bold = True
        hdr_cells[2].paragraphs[0].add_run(u'总评论数').bold = True
        hdr_cells[3].paragraphs[0].add_run(u'总阅读数').bold = True
        hdr_cells[4].paragraphs[0].add_run(u'总点赞数').bold = True
        hdr_cells[5].paragraphs[0].add_run(u'总收藏数').bold = True
        hdr_cells[6].paragraphs[0].add_run(u'事件热度').bold = True
        t5_sql = """
                    SELECT original_total, forward_total, reply_total, read_total, like_total, collect_total, event_heat
                    FROM (SELECT * FROM event_statistics_copy ORDER BY crawl_datetime DESC) a
                    WHERE event_id = %d
                    LIMIT 1
                    """ % event_id
        t5_resultList = session.execute(t5_sql)
        for r in t5_resultList:
            cells = table5.rows[1].cells
            cells[0].text = str(r[0])
            cells[1].text = str(r[1])
            cells[2].text = str(r[2])
            cells[3].text = str(r[3])
            cells[4].text = str(r[4])
            cells[5].text = str(r[5])
            cells[6].text = str(r[6])
            break
        for names in items[::-1]:
            if 'ori_forward_' + str(event_id) + '.jpg' == names:
                ori_forward_path = self.image_path + "//" + names
                document.add_picture(ori_forward_path, width=Inches(7))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_page_break()

        # 第六章
        para6 = document.add_paragraph()
        para6.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run6 = para6.add_run(u'第六章 网民地理分布（全国）Top10')
        run6.font.size = Pt(16)
        run6.bold = True
        table6 = document.add_table(rows=1, cols=3, style='Table Grid')
        table6.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells = table6.rows[0].cells
        hdr_cells[0].width = Inches(1.0)
        hdr_cells[1].width = Inches(1.0)
        hdr_cells[2].width = Inches(1.0)
        hdr_cells[0].paragraphs[0].add_run(u'排名').bold = True
        hdr_cells[1].paragraphs[0].add_run(u'省份').bold = True
        hdr_cells[2].paragraphs[0].add_run(u'人数').bold = True
        t6_sql = u"""
                    SELECT region, count(*) AS number
                    FROM (
                    SELECT substring_index(substring_index(ip_address,';',2),';',-1) AS region
                    FROM text_select_comment
                    WHERE ip_address LIKE '中国%%' AND event_id = %d
                    ) AS a
                    GROUP BY region
                    ORDER BY number DESC
                    LIMIT 10
                """ % event_id
        t6_resultList = session.execute(t6_sql)
        for r, j in zip(t6_resultList, range(1, 11)):
            cells = table6.add_row().cells
            cells[0].width = Inches(1.0)
            cells[1].width = Inches(1.0)
            cells[2].width = Inches(1.0)
            cells[0].text = str(j)
            cells[1].text = r[0]
            cells[2].text = str(r[1])
        for names in items[::-1]:
            if 'province_' + str(event_id) + '.jpg' == names:
                province_path = self.image_path + "//" + names
                document.add_picture(province_path, width=Inches(7))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_page_break()

        # 第六章
        para7 = document.add_paragraph()
        para7.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run7 = para7.add_run(u'第七章 网民地理分布（省内）Top10')
        run7.font.size = Pt(16)
        run7.bold = True
        table7 = document.add_table(rows=1, cols=3, style='Table Grid')
        table7.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells = table7.rows[0].cells
        hdr_cells[0].width = Inches(1.0)
        hdr_cells[1].width = Inches(1.0)
        hdr_cells[2].width = Inches(1.0)
        hdr_cells[0].paragraphs[0].add_run(u'排名').bold = True
        hdr_cells[1].paragraphs[0].add_run(u'市').bold = True
        hdr_cells[2].paragraphs[0].add_run(u'人数').bold = True
        t7_sql = u"""
            SELECT city, count(*) AS number
            FROM (
            SELECT substring_index(ip_address,';',-1) AS city
            FROM text_select_comment
            WHERE ip_address LIKE '中国;广东;%%' AND event_id = %d
            ) AS a
            WHERE city <> ''
            GROUP BY city
            ORDER BY number DESC
            LIMIT 10
        """ % event_id
        t7_resultList = session.execute(t7_sql)
        for r, j in zip(t7_resultList, range(1, 11)):
            cells = table7.add_row().cells
            cells[0].width = Inches(1.0)
            cells[1].width = Inches(1.0)
            cells[2].width = Inches(1.0)
            cells[0].text = str(j)
            cells[1].text = r[0]
            cells[2].text = str(r[1])
        for names in items[::-1]:
            if 'city_' + str(event_id) + '.jpg' == names:
                city_path = self.image_path + "//" + names
                document.add_picture(city_path, width=Inches(7))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # event_name = u'中大限外令'
        report_name = u'《' + event_name + u'》报告' + self.d + '.docx'
        document.save(self.report_path + '//' + report_name)

if __name__ == '__main__':
    print "——————————start——————————"
    start = time.time()
    allEventReport().Controller()
    # onlyEventReport().Controller(1)
    # allEventReport().newDocx()
    # onlyEventReport().newDocx(1)
    # allEventReport().eventHeatChart(1)
    session.close()
    end = time.time()
    print '%.4f' % (end-start)